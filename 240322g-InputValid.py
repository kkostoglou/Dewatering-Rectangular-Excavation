# DESCRIPTION:
# This program Calculates the required pumping discharge quantity Q, to succeed the desirable water drawdown, modeling the excavation as one big well
# It then generates a Word document with the input values, formulas, and results of the calculations.

from tkinter import *
from tkinter import messagebox
import math
from docx import Document
import os

def validate_real(entry):
    try:
        value = float(entry.get())
        return value
    except ValueError:
        messagebox.showerror("Error", "Please enter a valid real number")
        return None

def submit():
    real_value0 = validate_real(entry_real0)
    real_value1 = validate_real(entry_real1)
    real_value2 = validate_real(entry_real2)
    real_value3 = validate_real(entry_real3)
    real_value4 = validate_real(entry_real4)
    real_value5 = validate_real(entry_real5)

    if any(val is None for val in [real_value0, real_value1, real_value2, real_value3, real_value4, real_value5]):
        return

    K, H, hd0, hwl, a, b = real_value0, real_value1, real_value2, real_value3, real_value4, real_value5

    # Required flow drawdown (m)
    hd = hd0 - hwl
    # Hydraulic head at maximum dewatering (m)
    hw = H - hd
    # Radius of influence of Well or Point Source (m)
    r1 = 3000 * (H - hw) * math.sqrt(K)
    # Equivalent radius of the well (m)
    rw = math.sqrt(a * b / math.pi)
    # Total Radius of influence of Well (m)
    R = r1 + rw
    # Pumping rate (for all pumps) (m3/s)
    Q = (math.pi * K * (H ** 2 - hw ** 2)) / math.log(R / rw)
    # Pumping rate (for all pumps) (l/s)
    Qls = round(Q * 1000, 4)

    # Create a new Word document
    document = Document()

    # Add program description
    document.add_heading('Program Description', level=1)
    document.add_paragraph("This program Calculates the required pumping discharge quantity Q, to succeed")
    document.add_paragraph("the desirable water drawdown, modeling the rectangular excavation as one big well")

    # Add input values, formulas, and results to the document
    document.add_heading('Input Values', level=1)
    document.add_paragraph(f'Permeability index (K): {K:.2e} m/s')  # Modified to scientific notation
    document.add_paragraph(f'Hydraulic head of the original water table (H): {H:.2f} m')
    document.add_paragraph(f'Excavation depth from surface-lower water table (hd0): {hd0:.2f} m')
    document.add_paragraph(f'Ground water table depth, from surface (hwl): {hwl:.2f} m')
    document.add_paragraph(f'Length of excavation area (a): {a:.2f} m')
    document.add_paragraph(f'Width of excavation area (b): {b:.2f} m')

    document.add_heading('Formulas and Operations', level=1)
    document.add_paragraph('Required flow drawdown (hd) = hd0 - hwl')
    document.add_paragraph(f'hd = {hd0:.2f} - {hwl:.2f} = {hd:.2f} m')

    document.add_paragraph('Hydraulic head at maximum dewatering (hw) = H - hd')
    document.add_paragraph(f'hw = {H:.2f} - {hd:.2f} = {hw:.2f} m')

    document.add_paragraph('Radius of influence of Well or Point Source (r1) = 3000 * (H - hw) * sqrt(K)')
    document.add_paragraph(f'r1 = 3000 * ({H:.2f} - {hw:.2f}) * sqrt({K:.2e}) = {r1:.2f} m')

    document.add_paragraph('Equivalent radius of the well (rw) = sqrt(a * b / pi)')
    document.add_paragraph(f'rw = sqrt({a:.2f} * {b:.2f} / pi) = {rw:.2f} m')

    document.add_paragraph('Total Radius of influence of Well (R) = r1 + rw')
    document.add_paragraph(f'R = {r1:.2f} + {rw:.2f} = {R:.2f} m')

    document.add_paragraph('Pumping rate (Q) = (pi * K * (H^2 - hw^2)) / ln(R / rw)')
    document.add_paragraph(f'Q = (pi * {K:.2e} * ({H:.2f}^2 - {hw:.2f}^2)) / ln({R:.2f} / {rw:.2f}) = {Q:.4f} m3/s')  # Modified to four digits

    document.add_paragraph('Pumping rate (Qls) = Q * 1000')
    document.add_paragraph(f'Qls = {Q:.4f} * 1000 = {Qls:.2f} l/s')  # Modified to two digits

    # Save the document
    output_file_path = 'output.docx'
    document.save(output_file_path)
    messagebox.showinfo("Success", f"Output saved to {output_file_path}")

    # Open the saved document
    os.system(f'start {output_file_path}')

# Create the main window
root = Tk()
root.title("DISCHARGE PUMPING QUANTITY Q - DATA INPUT")
# Real Entry
label_real = Label(root, text="Permeability index : K (m/s)")
label_real.grid(row=0, column=0, padx=10, pady=5, sticky="e")
entry_real0 = Entry(root)
entry_real0.grid(row=0, column=1, padx=10, pady=5)
entry_real0.insert(0, "1.90E-06")  # Initial value

# Real Entry
label_real = Label(root, text="Hydraulic head of the original water table : H (m)")
label_real.grid(row=1, column=0, padx=10, pady=5, sticky="e")
entry_real1 = Entry(root)
entry_real1.grid(row=1, column=1, padx=10, pady=5)
entry_real1.insert(0, "90")  # Initial value

# Real Entry
label_real = Label(root, text="Excavation depth from surface-lower water table : hd0 (m)")
label_real.grid(row=2, column=0, padx=10, pady=5, sticky="e")
entry_real2 = Entry(root)
entry_real2.grid(row=2, column=1, padx=10, pady=5)
entry_real2.insert(0, "7.6")  # Initial value

# Real Entry
label_real = Label(root, text="Ground water table depth, from surface : hwl (m)")
label_real.grid(row=3, column=0, padx=10, pady=5, sticky="e")
entry_real3 = Entry(root)
entry_real3.grid(row=3, column=1, padx=10, pady=5)
entry_real3.insert(0, "0.9")  # Initial value

# Real Entry
label_real = Label(root, text="Length of excavation area : a (m)")
label_real.grid(row=4, column=0, padx=10, pady=5, sticky="e")
entry_real4 = Entry(root)
entry_real4.grid(row=4, column=1, padx=10, pady=5)
entry_real4.insert(0, "60")  # Initial value

# Real Entry
label_real = Label(root, text="Width of excavation area: b (m)")
label_real.grid(row=5, column=0, padx=10, pady=5, sticky="e")
entry_real5 = Entry(root)
entry_real5.grid(row=5, column=1, padx=10, pady=5)
entry_real5.insert(0, "25")  # Initial value

# Submit Button
submit_button = Button(root, text="Submit", command=submit)
submit_button.grid(row=6, columnspan=2, pady=10)

root.mainloop()