import sys
import csv
import math
import os
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QAction, QLabel, QVBoxLayout, QHBoxLayout,
    QPushButton, QWidget, QFileDialog, QTextEdit, QLineEdit, QMessageBox,
    QGridLayout, QSizePolicy, QSpacerItem
)
#import pandas as pd
from docx import Document
from PyQt5.QtGui import QPainter, QPen, QBrush, QFont, QPolygonF
from PyQt5.QtCore import Qt, QTimer, QPointF
from docx.shared import Pt

class RectangleCircleWidget_RESULTS(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMinimumSize(500, 500)  # Set the minimum size of the widget
        #self.a = 100  # Initial width of the rectangle
        #self.b = 100  # Initial height of the rectangle
        #self.R = 50   # Initial radius of the circle

        # Flag to indicate whether to draw the rectangle
        self.draw_rectangle_flag = False
        # Flag to indicate whether to draw the rectangle
        self.draw_RectangleCircle_flag = False

        self.curve_points = []  # Store points for drawing the curve


    def paintEvent(self, event):
        painter = QPainter(self)
        
        if self.draw_RectangleCircle_flag:
            # Draw the rectangle and circle
            #self.draw_rectangle_and_circle(painter)

            #def draw_rectangle_and_circle(self, painter):
            # Drawing rectangle and circle implementation remains the same as before
            a = self.a
            b = self.b

            pen_rectangle = QPen(Qt.black, 2, Qt.SolidLine)
            painter.setPen(pen_rectangle)
            brush = QBrush(Qt.blue, Qt.Dense3Pattern)
            brush.setColor(brush.color().lighter(40))
            painter.setBrush(brush)

            # Calculate the center of the rectangle
            rect_center_x = (self.width() - a * 2) / 3
            rect_center_y = (self.height() - b * 2) / 3 + 100
            
            # Round to integers
            rect_center_x = int(rect_center_x)
            rect_center_y = int(rect_center_y)

            # Calculate the position and size of the rectangle
            rect_x = rect_center_x - self.a // 2
            rect_y = rect_center_y - self.b // 2
            rect_x=int(rect_x)
            rect_y=int(rect_y)
            a_scaled = int(self.a * 4)  # Scale the width by 4
            b_scaled = int(self.b * 4)  # Scale the height by 4

            # Draw the rectangle with dimensions
            painter.drawRect(rect_x, rect_y, a_scaled, b_scaled)

            # Set smaller font1 size for drawing dimensions
            font1 = painter.font()
            font1.setPointSize(6)  # Adjust the font size as needed
            painter.setFont(font1)

            # Draw dimensions on the sides of the rectangle
            painter.drawText(rect_x - 60, rect_y + b_scaled // 2, f"{self.b}")  # Left side
            painter.drawText(rect_x + a_scaled + 10, rect_y + b_scaled // 2, f"{self.b}")  # Right side
            painter.drawText(rect_x + a_scaled // 2, rect_y - 10, f"{self.a}")  # Top side
            painter.drawText(rect_x + a_scaled // 2, rect_y + b_scaled + 30, f"{self.a}")  # Bottom side

            # Calculate the position of the circle's center
            circle_center_x = int(rect_x + a_scaled/2)
            circle_center_y = int(rect_y + b_scaled/2)
            # Calculate the scaled radius of the circle
            circle_radius_scaled = int(self.R * 4)  # Scale the radius by 4

            # Draw the scaled circle with the same center as the rectangle
            pen_circle = QPen(Qt.black, 2, Qt.DashLine)  # Set pen style to dashed line for circle
            pen_circle.setDashPattern([5, 5])  # Set the dash pattern to have longer dashes
            painter.setPen(pen_circle)
            brush_circle = QBrush(Qt.blue, Qt.Dense3Pattern)
            brush_circle.setColor(brush_circle.color().lighter(40))  # Lighter color with 60% transparency
            painter.setBrush(brush_circle)
            painter.drawEllipse(circle_center_x - circle_radius_scaled, circle_center_y - circle_radius_scaled, circle_radius_scaled * 2, circle_radius_scaled * 2)

            # Calculate the angle for the radius line (in radians)
            angle_radians = math.radians(-30)  # Example angle of 30 degrees anti-clock wise
            # Calculate the endpoint of the radius line
            end_x = circle_center_x + circle_radius_scaled * math.cos(angle_radians)
            end_y = circle_center_y + circle_radius_scaled * math.sin(angle_radians)

            # Draw the radius line
            #painter.drawLine(circle_center_x, circle_center_y, circle_center_x + circle_radius_scaled, circle_center_y)
            painter.drawLine(circle_center_x, circle_center_y, int(end_x), int(end_y))

            # Calculate the position for the text
            text_x = circle_center_x + circle_radius_scaled // 2
            text_y = circle_center_y - 10

            ## Draw the text
            #painter.drawText(text_x, text_y, f"R = {self.R}")  # Text on the radius line
            
            # Draw the text aligned with the radius line
            radius_text = f"R = {self.R:.2f}"  # Text on the radius line with two decimal points
            text_width = painter.fontMetrics().width(radius_text)
            text_height = painter.fontMetrics().height()
            
            # Calculate the angle of the radius line
            angle_degrees = math.degrees(angle_radians)
            ## Save the painter state
            painter.save()

            # Translate the coordinate system to the text position
            painter.translate(text_x, text_y)

            # Rotate the coordinate system by the angle
            painter.rotate(angle_degrees)

            # Draw the text
            painter.drawText(int(-text_width / 2 + 30), int(-text_height / 2 - 60), text_width, text_height, Qt.AlignCenter, radius_text)

            # Restore the painter state
            painter.restore()
            # Draw the curve
            painter.setRenderHint(QPainter.Antialiasing)  # Enable anti-aliasing for smoother lines
            pen_curve = QPen(Qt.red, 2, Qt.SolidLine)  # Set pen style for the curve
            pen_curve.setWidth(10)
            #painter = QPainter()
            painter.setPen(pen_curve)
            if self.curve_points:
                # Translate the coordinate system to the text position
                #painter.translate(circle_center_x-200, circle_center_y-100)
                #painter.drawLine(circle_center_x-200, circle_center_y-100, int(end_x), int(end_y))
                # Translate the coordinate system to the text position
                #painter.translate(circle_center_x-200, circle_center_y-100)
                #painter.drawPolyline(*self.curve_points)
                # Convert the list of points to QPolygonF
                # Draw polyline
                polyline = QPolygonF([QPointF(*point) for point in self.curve_points])
                painter.drawPolyline(polyline)
            
        # End the painting process
        #painter.end()
        
        
    def update_curve(self, curve_points):
        self.curve_points=curve_points
        self.update()

    def update_dimensions(self, a, b, R):
        self.a = a
        self.b = b
        self.R = R
        self.update()

    # Method to set the flag to True when "Run Analysis" is pressed
    def set_draw_rectangle_flag(self, flag):
        self.draw_rectangle_flag = flag

# Method to set the flag to True when "Run Analysis" is pressed
    def set_draw_RectangleCircle_flag(self, flag):
        self.draw_RectangleCircle_flag = flag       

class KKLikeApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.results_text = ""  # Define results_text as an empty string initially
        self.save_action = None
        self.analysis_menu = None
        self.setWindowTitle("KK-Like Application")
        self.setGeometry(100, 100, 2000, 1200)  # Increased width and height of the main window
        self.setMinimumSize(1500, 1800)  # Set minimum size of the main window

        # Initialize attributes for rectangle dimensions
        self.a = 100  # Example value, change as needed
        self.b = 100  # Example value, change as needed
        self.R = 50   # Initial radius of the circle
 
        #self.curve_points = []  # Store points for drawing the curve

        self.diak1 = 0
        #!!!!!!!!!!
        #self.diak2 = 0
        
        # Initialize attributes for input fields
        self.entry_real0_K = QLineEdit()
        self.entry_real1_H = QLineEdit()
        self.entry_real2_hd0 = QLineEdit()
        self.entry_real3_hwl = QLineEdit()
        self.entry_real4_a = QLineEdit()
        self.entry_real5_b = QLineEdit()

        # Set initial values for QLineEdit widgets
        self.entry_real0_K.setText("1.90E-06")
        self.entry_real1_H.setText("90")
        self.entry_real2_hd0.setText("7.6")
        self.entry_real3_hwl.setText("0.9")
        self.entry_real4_a.setText("60")
        self.entry_real5_b.setText("25")

        # Initialize input widgets
        self.entry_real4_a = QLineEdit()  # Initialize entry_real4_a attribute
        self.entry_real4_a.setFixedWidth(250)  # Set fixed width
        self.entry_real4_a.setText("60")  # Initial value

        self.entry_real5_b = QLineEdit()  # Initialize entry_real5_b attribute
        self.entry_real5_b.setFixedWidth(250)  # Set fixed width
        self.entry_real5_b.setText("25")  # Initial value

        self.output = QTextEdit()
        size_policy = QSizePolicy(QSizePolicy.Preferred, QSizePolicy.Maximum)
        self.output.setSizePolicy(size_policy)
        self.output.setMaximumHeight(1800)  # Set maximum height for QTextEdit
        self.output.setFixedWidth(300)  # Set fixed width for QTextEdit

        # Initialize the draw_rectangle_flag attribute
        self.draw_rectangle_flag = False
        self.draw_RectangleCircle_flag = False

        # Initialize the save_action attribute
        self.save_action = None

        # Create a single instance of RectangleCircleWidget_RESULTS
        self.RectangleCircle_widget = RectangleCircleWidget_RESULTS()
        
        # Initialize data input layout
        #self.data_input_layout = QVBoxLayout()
        #!!!!!!!!!!
        #self.data_input_layout = None
        
        self.init_ui()


    def init_ui(self):
        """Initialize the user interface of the application."""
        self.setup_central_widget()
        self.setup_menu_bar()
        self.setup_custom_widget()
        self.setup_output_panel()
        self.setup_initial_state()
        self.show()

    def setup_central_widget(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        self.layout = QVBoxLayout(central_widget)
        # Create a vertical layout for the main content
        self.content_layout = QVBoxLayout()
        self.content_layout.setAlignment(Qt.AlignLeft)  # Align to the left
        #self.layout.addLayout(self.content_layout)
        self.setup_input_panel()
        
        
    def setup_input_panel(self):
        input_panel_layout = QVBoxLayout()  # Make input_panel_layout an attribute
        self.content_layout.addLayout(input_panel_layout)
        # Geotechnical Data Input
        self.data_input_layout = QGridLayout()  # Make data_input_layout an attribute
        input_panel_layout.addLayout(self.data_input_layout)

    def setup_menu_bar(self):
        # Menus
        menu_bar = self.menuBar()
        # Set font size for menu items
        font = menu_bar.font()
        font.setPointSize(10)
        menu_bar.setFont(font)
        self.setup_file_menu(menu_bar)
        self.setup_analysis_menu(menu_bar)

    def setup_file_menu(self, menu_bar):
        file_menu = menu_bar.addMenu('  File')
        new_action = QAction('New', self)
        new_action.triggered.connect(self.show_geotechnical_data_input)
        # Apply the font to the QAction
        new_action.setFont(menu_bar.font())
        open_action = QAction('Open', self)
        open_action.triggered.connect(self.open_project)
        open_action.setFont(menu_bar.font())
        self.save_action = QAction('Save', self)  # Initialize save_action attribute
        self.save_action.triggered.connect(self.save_project)
        # Apply the font to the QAction
        self.save_action.setFont(menu_bar.font())
        exit_action = QAction('Exit', self)
        exit_action.triggered.connect(self.close)
        exit_action.setFont(menu_bar.font())
        file_menu.addActions([new_action, open_action, self.save_action, exit_action])

    def setup_analysis_menu(self, menu_bar):
        analysis_menu = menu_bar.addMenu('    Analysis')
        run_analysis_action = QAction('Run', self)
        run_analysis_action.triggered.connect(self.submit)
        run_analysis_action.setFont(menu_bar.font())
        results_action = QAction('Results', self)
        results_action.triggered.connect(self.handle_results_button)
        results_action.setFont(menu_bar.font())
        analysis_menu.addActions([run_analysis_action, results_action])
        self.analysis_menu = analysis_menu

    def setup_custom_widget(self):
        #self.RectangleCircle_widget = RectangleCircleWidget_RESULTS()
        RectangleCircle_layout = QVBoxLayout()
        RectangleCircle_layout.addWidget(self.RectangleCircle_widget)
        self.content_layout.addLayout(RectangleCircle_layout)
        self.layout.addLayout(self.content_layout, stretch=1)
        
        
    def setup_output_panel(self):
        self.output_panel_layout = QVBoxLayout()  # Make output_panel_layout an attribute  
        #self.layout.addLayout(self.output_panel_layout, stretch=1)  # Add output panel layout with adjusted stretch
        self.content_layout.addLayout(self.output_panel_layout)  
        self.output = QTextEdit()  
        self.output_panel_layout.addWidget(self.output)
        self.show()
    
        
    def setup_initial_state(self):
        self.save_action.setEnabled(False)
        self.analysis_menu.setEnabled(False)


    def draw_rectangle(self, painter):
        a = 100  # Initial width of the rectangle
        b = 100  # Initial height of the rectangle

        # Get the dimensions from the input labels
        length = float(self.entry_real4_a.text())
        width = float(self.entry_real5_b.text())

        # Set pen properties
        pen = QPen(Qt.black, 2, Qt.SolidLine)
        painter.setPen(pen)

        # Set brush properties for filling the rectangle
        # brush = QBrush(Qt.SolidPattern)  # Corrected
        brush = QBrush(Qt.blue, Qt.Dense3Pattern)  # Set blue color for brush
        brush.setColor(brush.color().lighter(40))  # Set transparency to 60%
        painter.setBrush(brush)

        # Calculate the position to draw the rectangle (centered horizontally)
        rect_x = (self.width() - a * 2) / 3
        rect_y = (self.height() - b * 2) / 3 + 50

        # Convert position and size to integers
        rect_x = int(rect_x)
        rect_y = int(rect_y)
        a = int(length)
        b = int(width)

        # Calculate scaled dimensions
        scaled_a = int(a * 4)  # Scale the width by 4
        scaled_b = int(b * 4)  # Scale the height by 4

        # Draw scaled rectangle
        painter.drawRect(rect_x, rect_y, scaled_a, scaled_b)

        # Update the widget to refresh the display
        self.update()
        
    def update_rectangle_dimensions(self, a, b):
        self.a = a
        self.b = b
        self.update()

    def show_geotechnical_data_input(self):
        # Clear layout before adding widgets
        if self.data_input_layout is not None:
            for i in reversed(range(self.data_input_layout.count())):
                widget = self.data_input_layout.itemAt(i).widget()
                if widget is not None:
                    widget.setParent(None)
        else:
            print("Error: self.data_input_layout is None")
        # Add widgets
        line_edit_width = 300  # Adjust this value as needed
        #self.diak2=self.diak2+1
        # Create the data input widgets
        label_string0 = QLabel("SINGLE BIG WELL - DATA INPUT\n")  # Input label
        label_real0_K = QLabel("Permeability index : K (m/s)")  # Input label
        self.entry_real0_K = QLineEdit()
        self.entry_real0_K.setFixedWidth(line_edit_width)
        self.entry_real0_K.setText("1.90E-06")  # Initial value

        label_real1_H = QLabel("Hydraulic head of the original water table : H (m)")  # Input label
        self.entry_real1_H = QLineEdit()
        self.entry_real1_H.setFixedWidth(line_edit_width)
        self.entry_real1_H.setText("90")  # Initial value

        label_real2_hd0 = QLabel("Excavation depth from surface-lower water table : hd0 (m)")  # Input label
        self.entry_real2_hd0 = QLineEdit()
        self.entry_real2_hd0.setFixedWidth(line_edit_width)
        self.entry_real2_hd0.setText("7.6")  # Initial value

        label_real3_hwl = QLabel("Ground water table depth, from surface : hwl (m)")  # Input label
        self.entry_real3_hwl = QLineEdit()
        self.entry_real3_hwl.setFixedWidth(line_edit_width)
        self.entry_real3_hwl.setText("0.9")  # Initial value

        label_real4_a = QLabel("Length of excavation area : a (m)")  # Input label
        self.entry_real4_a = QLineEdit()
        self.entry_real4_a.setFixedWidth(line_edit_width)
        self.entry_real4_a.setText("60")  # Initial value

        label_real5_b = QLabel("Width of excavation area: b (m)")  # Input label
        self.entry_real5_b = QLineEdit()
        self.entry_real5_b.setFixedWidth(line_edit_width)
        self.entry_real5_b.setText("25")  # Initial value

        # Add labels and input gadgets to layout
        self.data_input_layout.addWidget(label_string0, 0, 0)
        self.data_input_layout.addWidget(label_real0_K, 1, 0)
        self.data_input_layout.addWidget(label_real1_H, 2, 0)
        self.data_input_layout.addWidget(label_real2_hd0, 3, 0)
        self.data_input_layout.addWidget(label_real3_hwl, 4, 0)
        self.data_input_layout.addWidget(label_real4_a, 5, 0)
        self.data_input_layout.addWidget(label_real5_b, 6, 0)

        self.data_input_layout.addWidget(self.entry_real0_K, 1, 1, alignment=Qt.AlignLeft)
        self.data_input_layout.addWidget(self.entry_real1_H, 2, 1, alignment=Qt.AlignLeft)
        self.data_input_layout.addWidget(self.entry_real2_hd0, 3, 1, alignment=Qt.AlignLeft)
        self.data_input_layout.addWidget(self.entry_real3_hwl, 4, 1, alignment=Qt.AlignLeft)
        self.data_input_layout.addWidget(self.entry_real4_a, 5, 1, alignment=Qt.AlignLeft)
        self.data_input_layout.addWidget(self.entry_real5_b, 6, 1, alignment=Qt.AlignLeft)

        # Set vertical spacing between rows in the layout
        self.data_input_layout.setSpacing(8)  # Set vertical spacing to 10 pixels

        a = float(self.entry_real4_a.text())
        b = float(self.entry_real5_b.text())
        R = float(50)

        # Update dimensions of the rectangle widgetd
        self.RectangleCircle_widget.update_dimensions(a, b, R)

        # Set flag to True to draw the rectangle
        self.draw_rectangle_flag = True

        # When 'New Project' is pressed, enable the Save Project action,
        self.save_action.setEnabled(True)
        self.analysis_menu.setEnabled(True)

    def start_tkinter_event_loop(self):
        # Start Tkinter event loop
        self.tk_window = Tk()
        self.tk_window.title("SINGLE BIG WELL1 - DATA INPUT")
        self.tk_window.protocol("WM_DELETE_WINDOW", self.close_tk_window)  # Handle window close event
        self.tk_window.mainloop()  # Start Tkinter event loop

    def paintEvent(self, event):
        try:
            super().paintEvent(event)
            # Create painter object
            painter = QPainter(self)
            if self.draw_rectangle_flag:
                # Call draw_rectangle method with desired dimensions
                self.draw_rectangle(painter)
                # Draw dimensions on the sides of the rectangle
                a = float(self.entry_real4_a.text())
                b = float(self.entry_real5_b.text())
                rect_x = (self.width() - a * 2) / 3
                rect_y = (self.height() - b * 2) / 3 + 100
                rect_x = int(rect_x)
                rect_y = int(rect_y)
                a_scaled = int(a * 4)
                b_scaled = int(b * 4)
                # Calculate the positions for the text
                left_text_x = rect_x - 80
                left_text_y = rect_y + b_scaled // 2 - 90
                right_text_x = rect_x + a_scaled - 20
                right_text_y = rect_y + b_scaled // 2 - 90
                top_text_x = rect_x + a_scaled // 2 - 40
                top_text_y = rect_y - 110
                bottom_text_x = rect_x + a_scaled // 2 - 40
                bottom_text_y = rect_y + b_scaled - 80
 
                # Set smaller font1 size for drawing dimensions
                font1 = painter.font()
                font1.setPointSize(6)  # Adjust the font size as needed
                painter.setFont(font1)
            
                # Draw dimensions text
                painter.drawText(left_text_x, left_text_y, f"{b}")  # Left side
                painter.drawText(right_text_x, right_text_y, f"{b}")  # Right side
                painter.drawText(top_text_x, top_text_y, f"{a}")  # Top side
                painter.drawText(bottom_text_x, bottom_text_y, f"{a}")  # Bottom side
            # End the painting process
            #painter.end()
        except Exception as e:
            print("Error during painting:", e)

    def close_tk_window(self):
        self.tk_window.quit()
        self.tk_window.destroy()

    def process_tk_events(self):
        self.tk_window.update()
        QTimer.singleShot(10, self.process_tk_events)

    def submit(self):
        # Retrieve input values
        K = float(self.entry_real0_K.text())
        H = float(self.entry_real1_H.text())
        hd0 = float(self.entry_real2_hd0.text())
        hwl = float(self.entry_real3_hwl.text())
        a = float(self.entry_real4_a.text())
        b = float(self.entry_real5_b.text())

        # Perform calculations
        hd = hd0 - hwl  # Calculation: hd = hd0 - hwl
        hw = H - hd  # Calculation: hw = H - hd
        r1 = 3000 * (H - hw) * math.sqrt(K)  # Calculation: r1 = 3000 * (H - hw) * sqrt(K)
        rw = math.sqrt(a * b / math.pi)  # Calculation: rw = sqrt(a * b / pi)
        R = r1 + rw  # Calculation: R = r1 + rw
        Q = (math.pi * K * (H ** 2 - hw ** 2)) / math.log(R / rw)  # Calculation: Q = (pi * K * (H^2 - hw^2)) / ln(R / rw)
        Qls = round(Q * 1000, 4)  # Calculation: Qls = Q * 1000

        x_values = []
        y_values = []

        start_x = 1
        end_x = R
        delta_x = 1
        x = start_x
        while x <= end_x:
            y = math.sqrt(H ** 2 - (Q/(math.pi * K) * math.log(R / x)))
            #print(f"x = {x}, y = {y}")
            x_values.append(x)
            y_values.append(y)
            x += delta_x
        self.curve_points = [(int(x), int(y)) for x, y in zip(x_values, y_values)]
        self.RectangleCircle_widget.update_curve(self.curve_points)
            
        # Display results
        self.results_text = f"ANALYSIS RESULTS \n\n"
        self.results_text += f"Required flow drawdown (hd): {hd:.2f} m\n"
        self.results_text += f"Hydraulic head at maximum dewatering {hd:.2f} m\n"
        self.results_text += f"Radius of influence of Well or Point Source (r1): {r1:.2f} m\n"
        self.results_text += f"Equivalent radius of the well (rw): {rw:.2f} m\n"
        self.results_text += f"Total Radius of influence of Well (R): {R:.2f} m\n"
        self.results_text += f"Pumping rate (Q): {Q:.4f} m3\n"
        # Display results in the QTextEdit widget
        if not self.output:
            self.output = QTextEdit()
            self.output_panel_layout.addWidget(self.output)
        self.output.setText(self.results_text)

        # Save numerical data to a CSV file
        # Display results in the QTextEdit widget
        if not hasattr(self, 'tk_window'):
            self.tk_window = None
        if not hasattr(self, 'output'):
            self.output = None
        if not self.output:
            self.output = QTextEdit()
            self.output_panel_layout.addWidget(self.output)
        self.output.setText(self.results_text)

        data = {
            "K": K,
            "H": H,
            "hd0": hd0,
            "hwl": hwl,
            "a": a,
            "b": b,
            "hd": hd,
            "hw": hw,
            "r1": r1,
            "rw": rw,
            "R": R,
            "Q": Q,
            "Qls": Qls
        }
        # Write data to CSV file
        #df = pd.DataFrame(data, index=[0])
        #df.to_csv("project_data.csv", index=False)
        with open("project_data.csv", "w") as f:
            f.write("K,H,hd0,hwl,a,b,hd,hw,r1,rw,R,Q,Qls\n")
            for key in data:
                f.write(f"{data[key]},")
            f.write("\n")
        QMessageBox.information(self, "Success", "Analysis results and project data saved.")      
            
        # Create a new widget (e.g., RectangleCircleWidget_RESULTS)
        if  self.diak1 != 1:
            self.rectangle_circle_widget = RectangleCircleWidget_RESULTS()
            self.RectangleCircle_widget.update_curve(self.curve_points)
        if self.rectangle_circle_widget:  # Check if the widget is not None
            self.output_panel_layout.addWidget(self.rectangle_circle_widget)
            self.RectangleCircle_widget.update_curve(self.curve_points)
            self.diak1 = 1
     
        # Calculate the dimensions 'a' and 'b'
        a = float(self.entry_real4_a.text())
        b = float(self.entry_real5_b.text())
        
        # Set flag to True to draw the rectangle in RectangleWidget
        self.rectangle_circle_widget.set_draw_RectangleCircle_flag(True)

        # Update dimensions
        self.rectangle_circle_widget.update_dimensions(a, b, R)
        
        # Update dimensions
        self.RectangleCircle_widget.update_dimensions(a, b, R)

        # Update rectangle dimensions
        self.update_rectangle_dimensions(a, b)

        # Draw the rectangle in RectangleWidget before saving the results
        self.RectangleCircle_widget.repaint()

    def open_project(self):
        try:
            options = QFileDialog.Options()
            file_name, _ = QFileDialog.getOpenFileName(self, "Open Project", "", "CSV Files (*.csv);;All Files (*)", options=options)
            if file_name:
                ## Read project data from the CSV file
                #df = pd.read_csv(file_name)
                with open(file_name, 'r') as f:
                    reader = csv.DictReader(f)
                    data = next(reader)
                    # Set values to the respective line edits
                    self.entry_real0_K.setText(str(df['K'].values[0]))
                    self.entry_real1_H.setText(str(df['H'].values[0]))
                    self.entry_real2_hd0.setText(str(df['hd0'].values[0]))
                    self.entry_real3_hwl.setText(str(df['hwl'].values[0]))
                    self.entry_real4_a.setText(str(df['a'].values[0]))
                    self.entry_real5_b.setText(str(df['b'].values[0]))

                QMessageBox.information(self, "Opened Project", "Project data loaded successfully.")
                # When 'Open Project' is pressed, enable the Save Project action
                self.show_geotechnical_data_input()
                self.save_action.setEnabled(True)
                self.analysis_menu.setEnabled(True)
        except Exception as e:
            print("Error opening project:", e)

    def save_project(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "Save Project", "", "CSV Files (*.csv);;All Files (*)", options=options)
        if file_name:
            # Gather project data
            K = self.entry_real0_K.text()
            H = self.entry_real1_H.text()
            hd0 = self.entry_real2_hd0.text()
            hwl = self.entry_real3_hwl.text()
            a = self.entry_real4_a.text()
            b = self.entry_real5_b.text()

            # Write numerical project data to the CSV file
            data = {
                "K": float(K),
                "H": float(H),
                "hd0": float(hd0),
                "hwl": float(hwl),
                "a": float(a),
                "b": float(b)
            }
            #df = pd.DataFrame(data, index=[0])
            #if file_name:
            #    try:
            #        df.to_csv(file_name, index=False)
            #        QMessageBox.information(self, "Success", f"Project data saved to {file_name}")
            #    except Exception as e:
            #        QMessageBox.critical(self, "Error", f"An error occurred while saving the project: {str(e)}")

            fieldnames = data[0].keys()
            with open(file_name, 'w', newline='') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(data)

            QMessageBox.information(self, "Success", f"Project data saved to {file_name}")


        # After saving project, disable the Save Project action again
        self.save_action.setEnabled(False)

    def handle_results_button(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "Save Results", "", "Word Files (*.docx);;All Files (*)", options=options)
        if file_name:
            try:
                # Retrieve input values
                K = float(self.entry_real0_K.text())
                H = float(self.entry_real1_H.text())
                hd0 = float(self.entry_real2_hd0.text())
                hwl = float(self.entry_real3_hwl.text())
                a = float(self.entry_real4_a.text())
                b = float(self.entry_real5_b.text())

                results_text = self.results_text  # Assuming results_text is defined somewhere in your class
                hd = hd0 - hwl
                hw = H - hd
                r1 = 3000 * (H - hw) * math.sqrt(K)
                rw = math.sqrt(a * b / math.pi)
                R = r1 + rw
                Q = (math.pi * K * (H**2 - hw**2)) / math.log(R / rw)
                Qls = Q * 1000

                self.save_results_to_word(file_name, K, H, hd0, hwl, a, b, results_text, hd, hw, r1, rw, R, Q, Qls)

            except Exception as e:
                QMessageBox.critical(self, "Error", f"An error occurred while saving the results: {str(e)}")


    def save_results_to_word(self, file_name, K, H, hd0, hwl, a, b, results_text, hd, hw, r1, rw, R, Q, Qls):
        try:
            # Save results to a Word document
            doc = Document()
            # Set font size and line spacing
            font_size = Pt(10)  # Change the font size as needed
            line_spacing = Pt(12)  # Change the line spacing as needed
            for paragraph in doc.paragraphs:
                paragraph.style.font.size = font_size
                paragraph.paragraph_format.line_spacing = line_spacing
            doc.add_heading("Calculation of the required pumping discharge quantity Q, to succeed the desirable water drawdown, modeling the excavation as one big well", level=1)
            doc.add_heading('Input Data', level=1)
            input_text = f"Permeability index (K): {K:.2E} m/s\n"
            input_text += f"Hydraulic head of the original water table (H): {H} m\n"
            input_text += f"Excavation depth from surface-lower water table (hd0): {hd0} m\n"
            input_text += f"Ground water table depth, from surface (hwl): {hwl} m\n"
            input_text += f"Length of excavation area (a): {a} m\n"
            input_text += f"Width of excavation area (b): {b} m\n"
            doc.add_paragraph(input_text)

            doc.add_heading('Output Results', level=1)
            doc.add_paragraph(results_text)

            doc.add_heading('Formulas', level=1)
            doc.add_paragraph("hd = hd0 - hwl\n"
                              "hw = H - hd\n"
                              "r1 = 3000 * (H - hw) * sqrt(K)\n"
                              "rw = sqrt(a * b / pi)\n"
                              "R = r1 + rw\n"
                              "Q = (pi * K * (H^2 - hw^2)) / ln(R / rw)\n"
                              "Qls = Q * 1000")

            doc.add_heading('Symbolic Number Operations', level=1)

            doc.add_paragraph(
                "hd = hd0 - hwl\n"
                f"hd = {hd0} - {hwl}\n"
                f"hd = {hd:.2f} m\n\n"
                "hw = H - hd\n"
                f"hw = {H} - {hd:.2f}\n"
                f"hw = {hw:.2f} m\n\n"
                "r1 = 3000 * (H - hw) * sqrt(K)\n"
                f"r1 = 3000 * ({H} - {hw:.2f}) * sqrt({K})\n"
                f"r1 = {r1:.2f} m\n\n"
                "rw = sqrt(a * b / pi)\n"
                f"rw = sqrt({a} * {b} / pi)\n"
                f"rw = {rw:.2f} m\n\n"
                "R = r1 + rw\n"
                f"R = {r1:.2f} + {rw:.2f}\n"
                f"R = {R:.2f} m\n\n"
                "Q = (pi * K * (H^2 - hw^2)) / ln(R / rw)\n"
                f"Q = (pi * {K} * ({H}^2 - {hw:.2f}^2)) / ln({R:.2f} / {rw:.2f})\n"
                f"Q = {Q:.4f} m3\n\n"
                "Qls = Q * 1000\n"
                f"Qls = {Q:.4f} * 1000\n"
                f"Qls = {Qls} m3"
            )
            doc.save(file_name)

            # Open the saved Word file
            os.startfile(file_name)

            # Close the Tkinter window
            if hasattr(self, 'tk_window') and self.tk_window:
                self.tk_window.destroy()  # Close the Tkinter window

        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred while saving the results: {str(e)}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = KKLikeApp()
    window.show()
    print("Application started")
    sys.exit(app.exec())